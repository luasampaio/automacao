import docx
from docx import Document 


meu_word = Document()
meu_word.add_heading('Projeto de IA', 0)
    

meu_word.add_paragraph(
    'Itens: ', style='List Bullet'
)
meu_word.save('ia.docx')